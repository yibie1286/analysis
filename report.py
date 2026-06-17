"""
report.py — Production-level reporting system for Water Works Corporation CSI Survey.
Generates bilingual (English / Amharic) Word documents with charts, VIF, CI, and more.
"""

import io
import logging
from datetime import date

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


# ── Configuration ─────────────────────────────────────────────────────────────

class ReportConfig:
    BLUE   = RGBColor(0x1a, 0x6f, 0xa8)
    WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
    DARK   = RGBColor(0x1c, 0x2b, 0x3a)
    GREEN  = RGBColor(0x2e, 0x7d, 0x32)
    ORANGE = RGBColor(0xe6, 0x51, 0x00)
    GRAY   = RGBColor(0x6b, 0x7c, 0x93)
    RED    = RGBColor(0xc6, 0x28, 0x28)
    NYALA_FONT    = 'Nyala'
    BODY_SIZE     = 10.5
    TABLE_SIZE    = 9.5
    HEADING_COLOR = RGBColor(0x1a, 0x6f, 0xa8)


# Keep module-level aliases for backward compatibility
BLUE   = ReportConfig.BLUE
WHITE  = ReportConfig.WHITE
DARK   = ReportConfig.DARK
GREEN  = ReportConfig.GREEN
ORANGE = ReportConfig.ORANGE
GRAY   = ReportConfig.GRAY
RED    = ReportConfig.RED


# ── Amharic string table ──────────────────────────────────────────────────────

AM = {
    'title1':        'ውሃ ሥራዎች ኮርፖሬሽን',
    'title2':        'የደንበኛ እርካታ ጥናት',
    'title3':        'የትንተና ሪፖርት',
    'confidential':  'ሚስጥራዊ — ለውስጥ አጠቃቀም ብቻ',
    'date':          'ቀን',
    'respondents':   'ምላሽ ሰጪዎች',
    'overall_csi':   'አጠቃላይ CSI',
    'dimensions':    'ልኬቶች',
    'dim_sd':        'አገልግሎት አሰጣጥ · ቴክኒካዊ ጥራት · የፕሮጀክት አፈጻጸም · ግንኙነት',
    'exec_sum':      'አጭር ማጠቃለያ',
    'metric':        'መለኪያ',
    'value':         'ዋጋ',
    'interpretation':'ትርጓሜ',
    's1':            '1. መግቢያ',
    's1_body':       ('ውሃ ሥራዎች ኮርፖሬሽን (ውሥኮ) ለተቋማዊ ደንበኞቹ ከፍተኛ ጥራት ያለው '
                      'የውሃ መሠረተ ልማት አገልግሎት ለመስጠት ቁርጠኛ ነው። ይህ የደንበኛ እርካታ ጥናት '
                      'በዋና የአገልግሎት ልኬቶች ላይ የደንበኞችን አመለካከት ለመገምገም እና '
                      'ለቀጣይ ማሻሻያ ቦታዎችን ለመለየት ተዘጋጅቷል።'),
    's1_body2':      ('የጥናቱ መሣሪያ በKoboToolbox በኩል ተሰራጭቷል። ምላሾች '
                      'በ5-ነጥብ ሊከርት ሚዛን (1 = በጣም አልረካሁም፣ 5 = በጣም ረክቻለሁ) '
                      'ተሰብስበዋል።'),
    's1_1':          '1.1 የጥናቱ መዋቅር',
    's1_2':          '1.2 ናሙና',
    'dim_col':       'ልኬት',
    'vars_col':      'ተለዋዋጮች',
    'q_col':         'የጥያቄ ቁጥሮች',
    'items_col':     'የጥያቄ ብዛት',
    'scale_lbl':     'ሚዛን',
    'missing_lbl':   'የጎደሉ ዋጋዎች',
    'missing_val':   'በዓምድ አማካይ ተሞልቷል',
    's2':            '2. ገላጭ ስታቲስቲክስ',
    's2_body':       ('ገላጭ ስታቲስቲክስ ለእያንዳንዱ የጥናት ጥያቄ እና ልኬት '
                      'የምላሾቹን ማዕከላዊ አዝማሚያ እና ተለዋዋጭነት ያጠቃልላል።'),
    's2_1':          '2.1 በልኬት ደረጃ ማጠቃለያ',
    's2_2':          '2.2 በጥያቄ ደረጃ ስታቲስቲክስ',
    'dim_lbl':       'ልኬት',
    'mean_lbl':      'አማካይ',
    'std_lbl':       'መደበኛ ልዩነት',
    'item_lbl':      'ጥያቄ',
    'desc_lbl':      'መግለጫ',
    'min_lbl':       'ዝቅተኛ',
    'max_lbl':       'ከፍተኛ',
    'ci_95':         '95% እምነት ክልል',
    'rating_dist':      'የምላሽ ስርጭት',
    'rating_dist_body': 'ለእያንዳንዱ ልኬት በ1–5 ሚዛን የምላሾች ስርጭት ከዚህ በታች ቀርቧል።',
    'rating_1':         'በጣም አልረካሁም (1)',
    'rating_2':         'አልረካሁም (2)',
    'rating_3':         'በመጠኑ ረክቻለሁ (3)',
    'rating_4':         'ረክቻለሁ (4)',
    'rating_5':         'በጣም ረክቻለሁ (5)',
    'responses':        'ምላሾች',
    's3':            '3. የደንበኛ እርካታ መረጃ ጠቋሚ (CSI)',
    's3_body':       'CSI = (አማካይ ነጥብ ÷ 5) × 100 በሚለው ቀመር ይሰላል።',
    'interp_scale':  'የትርጓሜ ሚዛን',
    'csi_range':     'CSI ክልል',
    'very_sat':      'በጣም ረክቻለሁ',
    'satisfied':     'ረክቻለሁ',
    'neutral':       'በመጠኑ ረክቻለሁ',
    'dissatisfied':  'አልረካሁም',
    'overall_csi_lbl': 'አጠቃላይ CSI',
    's4':            '4. የአስተማማኝነት ትንተና — ክሮንባክ አልፋ',
    's4_body':       ('ክሮንባክ አልፋ (α) የእያንዳንዱ ልኬት ጥያቄዎች '
                      'የውስጥ ወጥነት ይለካል። α ≥ 0.7 ተቀባይነት ያለው ነው።'),
    's4_formula':    'ቀመር፡  α = (k / (k−1)) × (1 − Σσᵢ² / σ²_total)',
    'alpha_lbl':     'ክሮንባክ አልፋ',
    's5':            '5. ብዙ መስመራዊ ሪግሬሽን ትንተና',
    's5_body':       ('ብዙ መስመራዊ ሪግሬሽን የትኞቹ የአገልግሎት ልኬቶች '
                      'አጠቃላይ የደንበኛ እርካታን በከፍተኛ ሁኔታ እንደሚተነብዩ ለመለየት ጥቅም ላይ ውሏል።'),
    's5_model':      'SAT = β₀ + β₁·SD + β₂·TQ + β₃·PP + β₄·COM + ε',
    's5_1':          '5.1 የሞዴል ማጠቃለያ',
    's5_2':          '5.2 ቅንጅቶች',
    'sample_n':      'ናሙና መጠን (N)',
    'adj_r2':        'የተስተካከለ R²',
    'f_stat':        'F-ስታቲስቲክ',
    'model_p':       'የሞዴል p-ዋጋ',
    'model_sig':     'የሞዴል ጠቀሜታ',
    'sig_yes':       'ጠቃሚ (p < 0.05)',
    'sig_no':        'ጠቃሚ አይደለም',
    'variable':      'ተለዋዋጭ',
    'std_err':       'መደበኛ ስህተት',
    'significant':   'ጠቃሚ',
    'beta_strength': 'የቤታ ጥንካሬ',
    'strong':        'ጠንካራ',
    'moderate':      'መካከለኛ',
    'weak':          'ደካማ',
    's6':            '6. የትስስር ማትሪክስ',
    's6_body':       ('በልኬት ነጥቦች መካከል ያሉ የፒርሰን ትስስር ቅንጅቶች ከዚህ በታች ቀርበዋል። '
                      'ወደ 1.0 ቅርብ የሆኑ ዋጋዎች ጠንካራ አዎንታዊ ግንኙነት ያሳያሉ።'),
    's7_nps':        '7. የኔት ፕሮሞተር ነጥብ (NPS)',
    's7_nps_body':   ('NPS የደንበኛ ታማኝነት እና ኮርፖሬሽኑን ለሌሎች የመምከር ዕድልን ይለካል። '
                      'NPS = ((ፕሮሞተሮች − ዲትራክተሮች) / ጠቅላላ) × 100'),
    'promoters':     'ፕሮሞተሮች (9–10)',
    'passives':      'ፓሲቭ (7–8)',
    'detractors':    'ዲትራክተሮች (0–6)',
    'total':         'ጠቅላላ',
    'count':         'ቁጥር',
    'percentage':    'መቶኛ',
    'nps_score':     'NPS ነጥብ',
    'nps_excellent': 'እጅግ ጥሩ',
    'nps_good':      'ጥሩ',
    'nps_needs_imp': 'ማሻሻያ ያስፈልጋል',
    'nps_negative':  'አሉታዊ — ወዲያውኑ ይፈትሹ',
    'insights_title':'ዋና ዋና ግኝቶች',
    'insights_body': 'ከትንተናው ውጤቶች የተገኙ ዋና ዋና ግኝቶች፡',
    'recs_title':    'ምክረ ሃሳቦች',
    'recs_body':     ('የሚከተሉት ምክረ ሃሳቦች ከዚህ የጥናት ዑደት ትንተና ውጤቶች '
                      'በቀጥታ የተወሰዱ ናቸው፡'),
    'method_title':  'የዘዴ ማስታወሻ',
    'data_coll':     'የውሂብ ስብስብ',
    'data_coll_v':   'KoboToolbox የመስመር ላይ የጥናት መድረክ',
    'scale_m':       'ሚዛን',
    'scale_v':       '5-ነጥብ ሊከርት (1 = በጣም አልረካሁም፣ 5 = በጣም ረክቻለሁ)',
    'missing_m':     'የጎደሉ ዋጋዎች',
    'missing_mv':    'በዓምድ አማካይ ተሞልቷል',
    'csi_f':         'CSI ቀመር',
    'csi_fv':        'CSI (%) = (አማካይ ነጥብ / 5) × 100',
    'reliability_m': 'አስተማማኝነት',
    'reliability_v': 'ክሮንባክ አልፋ — የእያንዳንዱ ልኬት የውስጥ ወጥነት',
    'regression_m':  'ሪግሬሽን',
    'regression_v':  'OLS ብዙ መስመራዊ ሪግሬሽን',
    'corr_m':        'ትስስር',
    'corr_v':        'በልኬት ነጥቦች መካከል የፒርሰን ትስስር ቅንጅት',
    'software_m':    'ሶፍትዌር',
    'software_v':    'Python (pandas, numpy, statsmodels, scipy)',
    'gen_date':      'ሪፖርት የተዘጋጀበት ቀን',
    'footer':        'ውሃ ሥራዎች ኮርፖሬሽን — የደንበኛ እርካታ ሪፖርት — ሚስጥራዊ',
    'toc_title':     'የይዘት ዝርዝር',
    'priority_critical': '[ወሳኝ]',
    'priority_high':     '[ከፍተኛ]',
    'priority_medium':   '[መካከለኛ]',
    'priority_low':      '[ዝቅተኛ]',
    'rec_csi_critical':  'ወዲያውኑ ማሻሻያ ያስፈልጋል — CSI ከ60% በታች ነው። የምክንያት ትንተና ያካሂዱ።',
    'rec_csi_high':      'በደንበኛ ግብረ-መልስ ስብሰባዎች እና የሂደት ግምገማዎች ያሻሽሉ።',
    'rec_reg_top':       'በደንበኛ እርካታ ላይ ትልቁን ተጽዕኖ ያሳደረ ልኬት ነው — ቅድሚያ ይስጡት።',
    'rec_alpha_low':     'የጥያቄ ዝርዝሮቹን ይከልሱ — ከ0.7 በታች ያለ አልፋ ወጥ ያልሆነ መለኪያ ያሳያል።',
    'rec_nps_neg':       'የደንበኛ ታማኝነት ወዲያውኑ ይፈትሹ — NPS አሉታዊ ነው።',
    'rec_nps_low':       'የደንበኛ ታማኝነት ፕሮግራም ያዘጋጁ — NPS ዝቅተኛ ነው።',
    'rec_maintain':      'የአሁኑን የአገልግሎት ደረጃ ይጠብቁ እና ወቅታዊ ክትትል ያካሂዱ።',
    'rec_quarterly':     'ሩብ ዓመታዊ CSI ክትትል ዑደት ያቋቁሙ።',
    'rec_share':         'ግኝቶቹን ከክፍል ኃላፊዎች ጋር ያጋሩ እና ሊለኩ የሚችሉ KPI ያዘጋጁ።',
    # CSI interpretations
    'Very Satisfied':  'በጣም ረክቻለሁ',
    'Satisfied':       'ረክቻለሁ',
    'Neutral':         'በመጠኑ ረክቻለሁ',
    'Dissatisfied':    'አልረካሁም',
    # Alpha interpretations
    'Excellent':       'እጅግ ጥሩ',
    'Good':            'ጥሩ',
    'Acceptable':      'ተቀባይነት ያለው',
    'Questionable':    'አጠራጣሪ',
    'Poor':            'ደካማ',
    # Dimension names
    'Service Delivery':      'አገልግሎት አሰጣጥ',
    'Technical Quality':     'ቴክኒካዊ ጥራት',
    'Project Performance':   'የፕሮጀክት አፈጻጸም',
    'Communication':         'ግንኙነት',
    'Customer Satisfaction': 'የደንበኛ እርካታ',
}


# ── Validation ────────────────────────────────────────────────────────────────

def validate_result(result):
    """Raise ValueError if required keys are missing from result dict."""
    required = ['n', 'csi', 'overall_csi', 'overall_interp', 'reliability',
                'descriptive_items', 'descriptive_dims', 'insights']
    missing = [k for k in required if k not in result]
    if missing:
        raise ValueError(f"Result dict missing keys: {missing}")
    # Note:
    # 'rating_items' and 'rating_dims' are optional —
    # present only when generated by run_full_analysis() v2+.
    # Use result.get() to access them.
    # Do NOT add them to the required list.
    logger.info("Result validation passed (n=%s)", result.get('n'))


# ── Low-level helpers ─────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_font(cell, bold=False, color=None, size=None):
    for para in cell.paragraphs:
        for run in para.runs:
            if bold:  run.bold = bold
            if color: run.font.color.rgb = color
            if size:  run.font.size = Pt(size)


def _apply_amharic_font(run):
    """
    Set Nyala font on a run for correct Ethiopic rendering.
    FIX 2: Reuses existing w:rFonts element instead of appending duplicates.
    """
    run.font.name = ReportConfig.NYALA_FONT
    rPr = run._r.get_or_add_rPr()
    # Reuse existing rFonts element if already present — never append duplicates
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'),    ReportConfig.NYALA_FONT)
    rFonts.set(qn('w:hAnsi'),    ReportConfig.NYALA_FONT)
    rFonts.set(qn('w:cs'),       ReportConfig.NYALA_FONT)
    rFonts.set(qn('w:eastAsia'), ReportConfig.NYALA_FONT)


def add_page_break(doc):
    doc.add_page_break()


def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1a6fa8')
    pb.append(bottom)
    pPr.append(pb)


def styled_heading(doc, text, level=1, lang='en'):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = ReportConfig.BLUE
        if lang == 'am':
            _apply_amharic_font(run)
    return p


def body_text(doc, text, italic=False, color=None, lang='en'):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(ReportConfig.BODY_SIZE)
        run.font.color.rgb = color or DARK
        if italic: run.italic = True
        if lang == 'am':
            _apply_amharic_font(run)
    return p


def add_kv(doc, label, value, lang='en'):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f'{label}: ')
    r1.bold = True
    r1.font.size = Pt(ReportConfig.BODY_SIZE)
    r1.font.color.rgb = DARK
    if lang == 'am': _apply_amharic_font(r1)
    r2 = p.add_run(str(value))
    r2.font.size = Pt(ReportConfig.BODY_SIZE)
    r2.font.color.rgb = GRAY
    if lang == 'am': _apply_amharic_font(r2)


def add_styled_table(doc, headers, rows, lang='en'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], '1a6fa8')
        para = hdr[i].paragraphs[0]
        para.clear()
        run = para.add_run(str(h))
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(ReportConfig.TABLE_SIZE)
        if lang == 'am': _apply_amharic_font(run)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        tr = table.add_row()
        bg = 'f5f7fa' if ri % 2 == 0 else 'ffffff'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            para = cell.paragraphs[0]
            para.clear()
            run = para.add_run(str(val) if val is not None else '')
            run.font.size = Pt(ReportConfig.TABLE_SIZE)
            if lang == 'am': _apply_amharic_font(run)
            para.alignment = (WD_ALIGN_PARAGRAPH.LEFT
                              if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    return table


def safe_pct(part, total):
    """Safe percentage: returns 0.0 if total is zero."""
    return round(part / total * 100, 1) if total else 0.0


def t(key, lang, fallback=None):
    """Return Amharic string if lang=='am', else fallback or key itself."""
    if lang == 'am':
        return AM.get(key, fallback or key)
    return fallback or key


def csi_badge(val, lang='en'):
    if lang == 'am':
        if val >= 80: return f'{val}% ✦ {AM["very_sat"]}'
        if val >= 60: return f'{val}% ✔ {AM["satisfied"]}'
        if val >= 40: return f'{val}% ~ {AM["neutral"]}'
        return f'{val}% ✘ {AM["dissatisfied"]}'
    if val >= 80: return f'{val}% ✦ Very Satisfied'
    if val >= 60: return f'{val}% ✔ Satisfied'
    if val >= 40: return f'{val}% ~ Neutral'
    return f'{val}% ✘ Dissatisfied'


def alpha_badge(val, lang='en'):
    labels = {
        'en': {0.9: 'Excellent', 0.8: 'Good', 0.7: 'Acceptable',
               0.6: 'Questionable', 'low': 'Poor'},
        'am': {0.9: AM['Excellent'], 0.8: AM['Good'], 0.7: AM['Acceptable'],
               0.6: AM['Questionable'], 'low': AM['Poor']},
    }[lang]
    if val >= 0.9: return f'{val} — {labels[0.9]}'
    if val >= 0.8: return f'{val} — {labels[0.8]}'
    if val >= 0.7: return f'{val} — {labels[0.7]}'
    if val >= 0.6: return f'{val} — {labels[0.6]}'
    return f'{val} — {labels["low"]}'


def dim_name(name, lang):
    return AM.get(name, name) if lang == 'am' else name


def interp_name(name, lang):
    return AM.get(name, name) if lang == 'am' else name


def _beta_strength(beta, lang='en'):
    """Return a human-readable strength label for a regression beta coefficient."""
    ab = abs(beta)
    if ab >= 0.3: return t('strong',   lang, 'Strong')
    if ab >= 0.1: return t('moderate', lang, 'Moderate')
    return t('weak', lang, 'Weak')


def _nps_interp(score, lang='en'):
    """
    Return NPS interpretation label.
    FIX 6: Four tiers — Excellent / Good / Needs Improvement / Negative.
    """
    if score >= 50:
        return t('nps_excellent', lang, 'Excellent')
    if score >= 30:
        return t('nps_good', lang, 'Good')
    if score >= 0:
        return t('nps_needs_imp', lang, 'Needs Improvement')
    # Negative NPS — most urgent tier
    return AM.get('nps_negative', 'አሉታዊ — ወዲያውኑ ይፈትሹ') if lang == 'am' \
           else 'Negative — Address Urgently'


# ── Chart generation ──────────────────────────────────────────────────────────

def _get_amharic_font():
    """
    Find a system font that supports Ethiopic script for matplotlib.
    Returns a FontProperties object, or None to fall back to Latin labels.
    """
    from matplotlib.font_manager import FontProperties, findSystemFonts
    candidates = ['Nyala', 'Ebrima', 'Noto Sans Ethiopic',
                  'Abyssinica SIL', 'Ethiopia Jiret', 'GF Zemen Unicode']
    for name in candidates:
        try:
            fp = FontProperties(family=name)
            if fp.get_name() != 'DejaVu Sans':
                return fp
        except Exception:
            pass
    for fpath in findSystemFonts():
        fname = fpath.lower()
        if any(k in fname for k in ('nyala', 'ebrima', 'ethiopic',
                                     'noto', 'abyssinica')):
            try:
                return FontProperties(fname=fpath)
            except Exception:
                pass
    return None


def _safe_labels(labels, lang, font_prop):
    """
    If lang='am' and no Ethiopic font found, use Latin fallback labels
    so charts don't render as empty boxes.
    """
    if lang != 'am' or font_prop is not None:
        return labels
    fallbacks = {
        # Dimension names
        'አገልግሎት አሰጣጥ': 'SD',
        'ቴክኒካዊ ጥራት': 'TQ',
        'የፕሮጀክት አፈጻጸም': 'PP',
        'ግንኙነት': 'COM',
        'የደንበኛ እርካታ': 'SAT',
        # NPS labels — long form
        'ፕሮሞተሮች (9–10)': 'Promoters',
        'ፓሲቭ (7–8)': 'Passives',
        'ዲትራክተሮች (0–6)': 'Detractors',
        # NPS labels — short form
        'ፕሮሞተሮች': 'Promoters',
        'ፓሲቭ': 'Passives',
        'ዲትራክተሮች': 'Detractors',
        # Rating distribution labels
        '1-በጣም አልረካሁም': '1-Very Dissatisfied',
        '2-አልረካሁም': '2-Dissatisfied',
        '3-በመጠኑ ረክቻለሁ': '3-Neutral',
        '4-ረክቻለሁ': '4-Satisfied',
        '5-በጣም ረክቻለሁ': '5-Very Satisfied',
    }
    return [fallbacks.get(lbl, lbl) for lbl in labels]


def _make_chart(chart_type, labels, values, title, colors=None, lang='en'):
    """
    Generate a chart and return a seeked BytesIO buffer.

    Parameters
    ----------
    chart_type : str   'bar' or 'pie'
    labels     : list  Label strings
    values     : list  Numeric values
    title      : str   Chart title
    colors     : list  Optional hex/named colors
    lang       : str   'en' or 'am'

    Returns
    -------
    io.BytesIO  ready to pass to doc.add_picture()
    """
    default_colors = ['#1a6fa8', '#2e7d32', '#e65100', '#6b7c93', '#c62828']
    bar_colors = colors or default_colors[:len(labels)]

    am_font   = _get_amharic_font() if lang == 'am' else None
    safe_lbls = _safe_labels(labels, lang, am_font)

    def fkw(size=9):
        kw = {'fontsize': size}
        if lang == 'am' and am_font:
            kw['fontproperties'] = am_font
        return kw

    fig, ax = plt.subplots(figsize=(7, 4))

    if chart_type == 'bar':
        bars = ax.bar(safe_lbls, values, color=bar_colors,
                      edgecolor='white', linewidth=0.8)
        ax.set_ylabel('Value', fontsize=9)

        # FIX 3: Support both positive and negative beta values
        min_val = min(values) if values else 0
        max_val = max(values) if values else 1
        padding = max(abs(min_val), abs(max_val)) * 0.25 or 0.1
        ax.set_ylim(min_val - padding, max_val + padding)
        # Reference line at zero for charts with negative values
        if min_val < 0:
            ax.axhline(y=0, color='#cccccc', linewidth=0.8, linestyle='--')

        for bar, val in zip(bars, values):
            # Label above bar for positive, below for negative
            va  = 'bottom' if val >= 0 else 'top'
            off = padding * 0.15 if val >= 0 else -padding * 0.15
            ax.text(bar.get_x() + bar.get_width() / 2,
                    val + off,
                    f'{val:.3f}', ha='center', va=va, **fkw(8))

        ax.tick_params(axis='x', labelsize=8)
        ax.tick_params(axis='y', labelsize=8)
        if lang == 'am' and am_font:
            for tick in ax.get_xticklabels():
                tick.set_fontproperties(am_font)
                tick.set_fontsize(8)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

    elif chart_type == 'pie':
        wedge_colors = colors or ['#2e7d32', '#1a6fa8', '#c62828']
        wedges, texts, autotexts = ax.pie(
            values, labels=safe_lbls,
            colors=wedge_colors[:len(safe_lbls)],
            autopct='%1.1f%%', startangle=140,
            textprops=fkw(9),
        )
        if lang == 'am' and am_font:
            for txt in texts:
                txt.set_fontproperties(am_font)
                txt.set_fontsize(9)
        for at in autotexts:
            at.set_fontsize(8)

    title_kw = {'fontsize': 11, 'fontweight': 'bold', 'pad': 12}
    if lang == 'am' and am_font:
        title_kw['fontproperties'] = am_font
    ax.set_title(title, **title_kw)

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def _insert_chart(doc, chart_type, labels, values, title,
                  colors=None, lang='en'):
    """Wrap _make_chart in try/except so a chart failure never crashes the report."""
    try:
        buf = _make_chart(chart_type, labels, values, title,
                          colors=colors, lang=lang)
        doc.add_picture(buf, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        logger.info("Chart '%s' inserted successfully.", title)
    except Exception as exc:
        logger.warning("Chart generation failed for '%s': %s", title, exc)
        body_text(doc, f'[Chart unavailable: {exc}]', color=GRAY, lang=lang)


def _make_stacked_bar(dim_dist, lang='en', title=''):
    """
    Generate a stacked horizontal bar chart showing % of responses
    at each rating level (1-5) per dimension.

    Parameters
    ----------
    dim_dist : list
        result['rating_dims'] from analysis.py
    lang : str
        'en' or 'am'
    title : str
        Chart title

    Returns
    -------
    io.BytesIO
        PNG buffer ready for doc.add_picture()
    """
    if not dim_dist:
        return None

    # Color palette: red → orange → yellow → light green → dark green
    RATING_COLORS = ['#c62828', '#e65100', '#f9a825', '#66bb6a', '#2e7d32']

    RATING_LABELS_EN = [
        '1-Very Dissatisfied',
        '2-Dissatisfied',
        '3-Neutral',
        '4-Satisfied',
        '5-Very Satisfied',
    ]
    RATING_LABELS_AM = [
        '1-በጣም አልረካሁም',
        '2-አልረካሁም',
        '3-በመጠኑ ረክቻለሁ',
        '4-ረክቻለሁ',
        '5-በጣም ረክቻለሁ',
    ]
    rating_labels = RATING_LABELS_AM if lang == 'am' else RATING_LABELS_EN

    am_font = _get_amharic_font() if lang == 'am' else None

    # Build matrix — reverse so first dimension appears at top
    dim_labels  = []
    data_matrix = []
    for row in reversed(dim_dist):
        lbl = dim_name(row['Dimension'], lang)
        dim_labels.append(lbl)
        pcts = row.get('pcts', {})
        data_matrix.append([float(pcts.get(str(i), 0)) for i in range(1, 6)])

    n_dims = len(dim_labels)
    fig, ax = plt.subplots(figsize=(9, max(3.5, n_dims * 0.8 + 1.5)))

    y_pos = range(n_dims)
    lefts = [0.0] * n_dims
    bars_list = []

    for ci, (color, label) in enumerate(zip(RATING_COLORS, rating_labels)):
        vals = [data_matrix[ri][ci] for ri in range(n_dims)]
        bars = ax.barh(list(y_pos), vals, left=lefts,
                       color=color, label=label, height=0.55)
        bars_list.append((bars, vals))
        lefts = [lefts[ri] + vals[ri] for ri in range(n_dims)]

    # Labels inside bars — show if segment wide enough (≥5%), outside if 3–4%
    for bars, vals in bars_list:
        for bar, val in zip(bars, vals):
            if val <= 0:
                continue
            cx = bar.get_x() + bar.get_width() / 2
            cy = bar.get_y() + bar.get_height() / 2
            if val >= 5:
                ax.text(cx, cy, f'{val:.0f}%',
                        ha='center', va='center',
                        fontsize=7, color='white', fontweight='bold')
            elif val >= 2:
                # Place label just outside the right edge of the segment
                ax.text(bar.get_x() + bar.get_width() + 0.5, cy,
                        f'{val:.0f}%',
                        ha='left', va='center',
                        fontsize=6, color='#333333')

    ax.set_yticks(list(y_pos))
    safe_lbls = _safe_labels(dim_labels, lang, am_font)
    ax.set_yticklabels(safe_lbls, fontsize=9)
    if lang == 'am' and am_font:
        for tick in ax.get_yticklabels():
            tick.set_fontproperties(am_font)
            tick.set_fontsize(9)

    ax.set_xlabel('% of Responses', fontsize=9)
    ax.set_xlim(0, 105)   # slight extra room for outside labels
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

    # Legend below the chart — enough bottom margin so it doesn't overlap x-label
    legend_kw = {
        'fontsize': 8,
        'loc': 'upper center',
        'bbox_to_anchor': (0.5, -0.18),
        'ncol': 3,
        'frameon': False,
    }
    if lang == 'am' and am_font:
        legend_kw['prop'] = am_font
    ax.legend(rating_labels, **legend_kw)

    title_kw = {'fontsize': 11, 'fontweight': 'bold', 'pad': 10}
    if lang == 'am' and am_font:
        title_kw['fontproperties'] = am_font
    ax.set_title(title, **title_kw)

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def _insert_stacked_bar(doc, dim_dist, lang='en', title=''):
    """Insert stacked bar chart — fails gracefully if chart errors."""
    try:
        buf = _make_stacked_bar(dim_dist, lang=lang, title=title)
        if buf is None:
            return
        doc.add_picture(buf, width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        logger.info("Stacked bar chart '%s' inserted.", title)
    except Exception as exc:
        logger.warning("Stacked bar chart failed: %s", exc)
        body_text(doc, f'[Rating distribution chart unavailable: {exc}]',
                  color=GRAY, lang=lang)


# ── Header / Footer / TOC ─────────────────────────────────────────────────────

def _add_header(doc, org_name, lang='en'):
    """Add a real section header with the organisation name."""
    section = doc.sections[0]
    header  = section.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    run = para.add_run(org_name)
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run.italic = True
    if lang == 'am':
        _apply_amharic_font(run)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_footer(doc, lang='en'):
    """Add a footer with the confidential label and an auto page-number field."""
    section = doc.sections[0]
    footer  = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    run_left = para.add_run(t('footer', lang,
        'Water Works Corporation — Customer Satisfaction Report — CONFIDENTIAL'))
    run_left.font.size = Pt(8)
    run_left.font.color.rgb = GRAY
    run_left.italic = True
    if lang == 'am':
        _apply_amharic_font(run_left)
    # Tab + page-number field
    run_tab = para.add_run('\t')
    run_tab.font.size = Pt(8)
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run_pg = para.add_run()
    run_pg._r.append(fldChar_begin)
    run_pg._r.append(instrText)
    run_pg._r.append(fldChar_end)
    run_pg.font.size = Pt(8)
    run_pg.font.color.rgb = GRAY
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_table_of_contents(doc, lang='en'):
    """Insert a Table of Contents placeholder using Word's TOC field."""
    styled_heading(doc, t('toc_title', lang, 'Table of Contents'),
                   level=1, lang=lang)
    add_divider(doc)
    para = doc.add_paragraph()
    run  = para.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_sep)
    run._r.append(fldChar_end)
    note = ('(Right-click and select "Update Field" to populate the Table of Contents.)'
            if lang == 'en' else
            '(ቀኝ ጠቅ ያድርጉ እና "Update Field" ይምረጡ።)')
    body_text(doc, note, italic=True, color=GRAY, lang=lang)
    add_page_break(doc)
    logger.info("Table of contents placeholder inserted.")


# ── Smart recommendations ─────────────────────────────────────────────────────

def dynamic_recommendations(result, lang='en'):
    """Generate prioritised recommendations from analysis results."""
    csi_map  = {r['Code']: r for r in result['csi']}
    rel_map  = {r['Code']: r for r in result['reliability']}
    reg      = result.get('regression')
    nps      = result.get('nps')

    sig_codes = set()
    if reg:
        sig_codes = {c['Code'] for c in reg['coefficients']
                     if c['Significant'] == 'Yes'}

    def _pfx(tier):
        return {
            'CRITICAL': t('priority_critical', lang, '[CRITICAL]'),
            'HIGH':     t('priority_high',     lang, '[HIGH]'),
            'MEDIUM':   t('priority_medium',   lang, '[MEDIUM]'),
            'LOW':      t('priority_low',      lang, '[LOW]'),
        }[tier]

    recs = []

    for code, row in csi_map.items():
        if code == 'SAT':
            continue
        dn      = dim_name(row['Dimension'], lang)
        csi_val = row['CSI (%)']
        if csi_val < 60 and code in sig_codes:
            tier = 'CRITICAL'
            msg  = t('rec_csi_critical', lang,
                     'Urgently address — CSI below 60% and statistically significant. '
                     'Conduct root-cause analysis.')
        elif csi_val < 60:
            tier = 'HIGH'
            msg  = t('rec_csi_critical', lang,
                     'Urgently address — CSI falls below Satisfied threshold. '
                     'Conduct root-cause analysis.')
        elif csi_val < 75:
            tier = 'HIGH'
            msg  = t('rec_csi_high', lang,
                     'Improve through structured client feedback sessions and process reviews.')
        else:
            continue
        recs.append(f'{_pfx(tier)} {dn} (CSI {csi_val}%): {msg}')

    if reg:
        sig = [c for c in reg['coefficients'] if c['Significant'] == 'Yes']
        if sig:
            top  = max(sig, key=lambda x: abs(x['Beta']))
            dn   = dim_name(top['Variable'], lang)
            tier = ('CRITICAL'
                    if csi_map.get(top.get('Code', ''), {}).get('CSI (%)', 100) < 60
                    else 'HIGH')
            msg  = t('rec_reg_top', lang,
                     'Strongest significant predictor of satisfaction — '
                     'prioritize investment here.')
            recs.append(f'{_pfx(tier)} {dn} (β = {top["Beta"]}): {msg}')

    for code, row in rel_map.items():
        if row['Cronbach Alpha'] < 0.7:
            dn  = dim_name(row['Dimension'], lang)
            msg = t('rec_alpha_low', lang,
                    'Review questionnaire items — Alpha below 0.7 indicates '
                    'inconsistent measurement.')
            recs.append(f'{_pfx("HIGH")} {dn} (α = {row["Cronbach Alpha"]}): {msg}')

    if nps:
        if nps['nps_score'] < 0:
            msg = t('rec_nps_neg', lang,
                    f"Address client loyalty urgently — NPS of {nps['nps_score']} is negative.")
            recs.append(f'{_pfx("MEDIUM")} {msg}')
        elif nps['nps_score'] < 30:
            msg = t('rec_nps_low', lang,
                    f"Develop a client loyalty program — NPS of {nps['nps_score']} "
                    f"indicates room for improvement.")
            recs.append(f'{_pfx("MEDIUM")} {msg}')

    if not recs:
        recs.append(
            f'{_pfx("LOW")} '
            f'{t("rec_maintain", lang, "Maintain current service standards and conduct periodic follow-up surveys.")}'
        )

    recs.append(
        f'{_pfx("LOW")} '
        f'{t("rec_quarterly", lang, "Establish a quarterly CSI monitoring cycle to track improvement progress.")}'
    )
    recs.append(
        f'{_pfx("LOW")} '
        f'{t("rec_share", lang, "Share findings with department heads and develop dimension-specific action plans with measurable KPIs.")}'
    )
    logger.info("Generated %d recommendations.", len(recs))
    return recs


# ── Section builders ──────────────────────────────────────────────────────────

def add_cover_page(doc, result, lang, cfg):
    """Render the cover page."""
    logger.info("Building cover page (lang=%s).", lang)
    for _ in range(3):
        doc.add_paragraph()

    if lang == 'am':
        org_line = 'ውሃ ሥራዎች ኮርፖሬሽን'
        sub_line = 'WATER WORKS CORPORATION'
        t1_line  = 'የደንበኛ እርካታ ጥናት'
        t2_line  = 'የትንተና ሪፖርት'
    else:
        org_line = 'WATER WORKS CORPORATION'
        sub_line = 'ውሃ ሥራዎች ኮርፖሬሽን'
        t1_line  = 'Customer Satisfaction Survey'
        t2_line  = 'Analysis Report'

    for txt, sz, bold, color in [
        (org_line, 16, True,  cfg.BLUE),
        (sub_line, 12, False, cfg.GRAY),
        ('',        6, False, cfg.GRAY),
        (t1_line,  26, True,  cfg.DARK),
        (t2_line,  22, True,  cfg.BLUE),
    ]:
        p = doc.add_paragraph(txt)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(sz)
            run.bold = bold
            run.font.color.rgb = color
            if lang == 'am':
                _apply_amharic_font(run)

    doc.add_paragraph()
    doc.add_paragraph()
    add_divider(doc)
    doc.add_paragraph()

    overall_csi    = result['overall_csi']
    overall_interp = result['overall_interp']
    n              = result['n']
    interp_display = interp_name(overall_interp, lang)

    meta = [
        (t('date',        lang, 'Date'),         date.today().strftime('%B %d, %Y')),
        (t('respondents', lang, 'Respondents'),   str(n)),
        (t('overall_csi', lang, 'Overall CSI'),   f'{overall_csi}% — {interp_display}'),
        (t('dimensions',  lang, 'Dimensions'),
         t('dim_sd', lang,
           'Service Delivery · Technical Quality · Project Performance · Communication')),
    ]
    for label, val in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f'{label}:  ')
        r1.bold = True
        r1.font.color.rgb = cfg.DARK
        r2 = p.add_run(val)
        r2.font.color.rgb = cfg.GRAY
        if lang == 'am':
            _apply_amharic_font(r1)
            _apply_amharic_font(r2)

    doc.add_paragraph()
    conf = doc.add_paragraph(t('confidential', lang,
                               'CONFIDENTIAL — For Internal Use Only'))
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in conf.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = cfg.GRAY
        run.italic = True
        if lang == 'am':
            _apply_amharic_font(run)

    add_page_break(doc)


def add_executive_summary(doc, result, lang, cfg):
    """Render the Executive Summary section."""
    logger.info("Building executive summary.")
    styled_heading(doc, t('exec_sum', lang, 'Executive Summary'), lang=lang)
    add_divider(doc)
    doc.add_paragraph()

    overall_csi    = result['overall_csi']
    overall_interp = result['overall_interp']
    n              = result['n']
    reg            = result.get('regression')
    nps            = result.get('nps')
    interp_display = interp_name(overall_interp, lang)

    csi_non_sat = [r for r in result['csi'] if r['Code'] != 'SAT']
    best_dim  = max(csi_non_sat, key=lambda x: x['CSI (%)']) if csi_non_sat else None
    worst_dim = min(csi_non_sat, key=lambda x: x['CSI (%)']) if csi_non_sat else None

    if lang == 'am':
        summary = (f'ይህ ሪፖርት ከ{n} ተቋማዊ ደንበኞች የተሰበሰበ የደንበኛ እርካታ ጥናት ውጤቶችን ያቀርባል። '
                   f'አጠቃላይ CSI {overall_csi}% ሲሆን ደንበኞቹ "{interp_display}" ናቸው። ')
        if best_dim:
            summary += (f'{dim_name(best_dim["Dimension"], lang)} '
                        f'ከፍተኛ CSI {best_dim["CSI (%)"]}% አስመዝግቧል። ')
        if worst_dim:
            summary += (f'{dim_name(worst_dim["Dimension"], lang)} '
                        f'ዝቅተኛ CSI {worst_dim["CSI (%)"]}% አስመዝግቧል። ')
    else:
        summary = (f'This report presents findings of a customer satisfaction survey among {n} '
                   f'institutional clients of Water Works Corporation. '
                   f'The Overall CSI stands at {overall_csi}%, indicating clients are '
                   f'generally "{interp_display}". ')
        if best_dim:
            summary += (f'{best_dim["Dimension"]} recorded the highest CSI '
                        f'at {best_dim["CSI (%)"]}%. ')
        if worst_dim:
            summary += f'{worst_dim["Dimension"]} recorded the lowest at {worst_dim["CSI (%)"]}%. '
        if reg:
            sig = [c for c in reg['coefficients'] if c['Significant'] == 'Yes']
            if sig:
                top = max(sig, key=lambda x: abs(x['Beta']))
                summary += (f'Regression identifies {top["Variable"]} as the strongest '
                            f'predictor (β = {top["Beta"]}, R² = {reg["r_squared"]}).')

    body_text(doc, summary, lang=lang)
    doc.add_paragraph()

    kpi_rows = [
        [t('overall_csi_lbl', lang, 'Overall CSI'), f'{overall_csi}%', interp_display],
    ]
    for row in result['csi']:
        kpi_rows.append([dim_name(row['Dimension'], lang),
                         f"{row['CSI (%)']}%",
                         interp_name(row['Interpretation'], lang)])
    if reg:
        kpi_rows.append(['R²', str(reg['r_squared']),
                         f"{round(reg['r_squared'] * 100, 1)}%"])
    if nps:
        kpi_rows.append([t('nps_score', lang, 'NPS Score'),
                         str(nps['nps_score']),
                         _nps_interp(nps['nps_score'], lang)])

    add_styled_table(doc,
        [t('metric', lang, 'Metric'),
         t('value',  lang, 'Value'),
         t('interpretation', lang, 'Interpretation')],
        kpi_rows, lang=lang)
    add_page_break(doc)


def add_descriptive_statistics(doc, result, lang, cfg):
    """Render Section 2 — Descriptive Statistics with 95% CI."""
    logger.info("Building descriptive statistics section.")
    from scipy import stats as scipy_stats

    styled_heading(doc, t('s2', lang, '2. Descriptive Statistics'), lang=lang)
    add_divider(doc)
    body_text(doc, t('s2_body', lang,
        'Descriptive statistics summarize the central tendency and variability '
        'of responses for each survey item and dimension.'), lang=lang)

    styled_heading(doc, t('s2_1', lang, '2.1 Dimension-Level Summary'),
                   level=2, lang=lang)
    dim_rows = [[dim_name(r['Dimension'], lang), r['Mean'], r['Std Dev']]
                for r in result['descriptive_dims']]
    add_styled_table(doc,
        [t('dim_lbl', lang, 'Dimension'),
         t('mean_lbl', lang, 'Mean'),
         t('std_lbl',  lang, 'Std Dev')],
        dim_rows, lang=lang)

    styled_heading(doc, t('s2_2', lang, '2.2 Item-Level Statistics'),
                   level=2, lang=lang)
    n = result['n']
    item_rows = []
    for r in result['descriptive_items']:
        mean = r['Mean']
        std  = r['Std Dev']
        try:
            se     = std / (n ** 0.5) if n > 1 else 0.0
            ci     = scipy_stats.t.interval(0.95, df=max(n - 1, 1),
                                            loc=mean, scale=se)
            ci_str = f'[{ci[0]:.2f}, {ci[1]:.2f}]'
        except Exception:
            ci_str = 'N/A'
        item_rows.append([r['Item'], r.get('Description', r['Item']),
                          mean, std, r['Min'], r['Max'], ci_str])
    add_styled_table(doc,
        [t('item_lbl', lang, 'Item'),
         t('desc_lbl', lang, 'Description'),
         t('mean_lbl', lang, 'Mean'),
         t('std_lbl',  lang, 'Std Dev'),
         t('min_lbl',  lang, 'Min'),
         t('max_lbl',  lang, 'Max'),
         t('ci_95',    lang, '95% CI')],
        item_rows, lang=lang)
    add_page_break(doc)


def add_rating_distribution(doc, result, lang, cfg):
    """
    Render rating distribution — stacked bar chart + summary table.
    Shows how responses spread across 1-5 satisfaction levels per dimension.
    """
    logger.info("Building rating distribution section.")
    dim_dist = result.get('rating_dims', [])
    if not dim_dist:
        logger.warning("No rating_dims data — skipping distribution section.")
        return

    styled_heading(doc,
        t('rating_dist', lang, '3.1 Response Rating Distribution'),
        level=2, lang=lang)
    body_text(doc,
        t('rating_dist_body', lang,
          'The chart below shows the distribution of responses '
          'across the 1–5 satisfaction scale for each dimension.'),
        lang=lang)

    chart_title = ('Response Distribution by Dimension (%)'
                   if lang == 'en' else 'የምላሽ ስርጭት በልኬት (%)')
    _insert_stacked_bar(doc, dim_dist, lang=lang, title=chart_title)

    rating_headers = [
        t('dim_lbl',      lang, 'Dimension'),
        t('responses',    lang, 'Responses'),
        '% ' + t('satisfied',    lang, 'Satisfied')    + ' (4–5)',
        '% ' + t('neutral',      lang, 'Neutral')      + ' (3)',
        '% ' + t('dissatisfied', lang, 'Dissatisfied') + ' (1–2)',
    ]
    rating_rows = []
    for row in dim_dist:
        pcts     = row.get('pcts', {})
        sat_pct  = round(float(pcts.get('4', 0)) + float(pcts.get('5', 0)), 1)
        neu_pct  = round(float(pcts.get('3', 0)), 1)
        dis_pct  = round(float(pcts.get('1', 0)) + float(pcts.get('2', 0)), 1)
        rating_rows.append([
            dim_name(row['Dimension'], lang),
            str(row.get('total', 0)),
            f'{sat_pct}%',
            f'{neu_pct}%',
            f'{dis_pct}%',
        ])
    add_styled_table(doc, rating_headers, rating_rows, lang=lang)


def add_csi_section(doc, result, lang, cfg):
    """Render Section 3 — CSI with bar chart."""
    logger.info("Building CSI section.")
    overall_csi    = result['overall_csi']
    overall_interp = result['overall_interp']
    interp_display = interp_name(overall_interp, lang)

    styled_heading(doc,
        t('s3', lang, '3. Customer Satisfaction Index (CSI)'), lang=lang)
    add_divider(doc)
    body_text(doc, t('s3_body', lang, 'CSI (%) = (Mean Score ÷ 5) × 100'),
              lang=lang)

    body_text(doc, t('interp_scale', lang, 'Interpretation scale:'),
              italic=True, lang=lang)
    interp_rows = [
        ['80–100%', t('very_sat',     lang, 'Very Satisfied')],
        ['60–79%',  t('satisfied',    lang, 'Satisfied')],
        ['40–59%',  t('neutral',      lang, 'Neutral')],
        ['< 40%',   t('dissatisfied', lang, 'Dissatisfied')],
    ]
    add_styled_table(doc,
        [t('csi_range',      lang, 'CSI Range'),
         t('interpretation', lang, 'Interpretation')],
        interp_rows, lang=lang)

    # FIX 2 (CSI): show plain percentage — no badge duplication
    csi_rows = [
        [dim_name(r['Dimension'], lang),
         r['Mean Score'],
         f"{r['CSI (%)']}%",
         interp_name(r['Interpretation'], lang)]
        for r in result['csi']
    ]
    add_styled_table(doc,
        [t('dim_lbl',  lang, 'Dimension'),
         t('mean_lbl', lang, 'Mean Score'),
         'CSI (%)',
         t('interpretation', lang, 'Interpretation')],
        csi_rows, lang=lang)

    # Rating distribution sub-section
    add_rating_distribution(doc, result, lang, cfg)

    # CSI bar chart
    non_sat = [r for r in result['csi'] if r['Code'] != 'SAT']
    if non_sat:
        chart_labels = [dim_name(r['Dimension'], lang) for r in non_sat]
        chart_values = [r['CSI (%)'] for r in non_sat]
        chart_title  = ('CSI by Dimension (%)' if lang == 'en'
                        else 'CSI በልኬት (%)')
        _insert_chart(doc, 'bar', chart_labels, chart_values,
                      chart_title, lang=lang)

    p  = doc.add_paragraph()
    r1 = p.add_run(f'{t("overall_csi_lbl", lang, "Overall CSI")}: ')
    r1.bold = True
    r1.font.size = Pt(12)
    if lang == 'am': _apply_amharic_font(r1)
    r2 = p.add_run(f'{overall_csi}%  —  {interp_display}')
    r2.font.size = Pt(12)
    r2.bold = True
    r2.font.color.rgb = cfg.GREEN if overall_csi >= 60 else cfg.ORANGE
    if lang == 'am': _apply_amharic_font(r2)
    doc.add_paragraph()
    add_page_break(doc)


def add_reliability_section(doc, result, lang, cfg):
    """Render Section 4 — Reliability Analysis."""
    logger.info("Building reliability section.")
    styled_heading(doc,
        t('s4', lang, "4. Reliability Analysis — Cronbach's Alpha"), lang=lang)
    add_divider(doc)
    body_text(doc, t('s4_body', lang,
        "Cronbach's Alpha (α) measures internal consistency. "
        "α ≥ 0.7 is acceptable."), lang=lang)
    body_text(doc, t('s4_formula', lang,
        'Formula: α = (k / (k−1)) × (1 − Σσᵢ² / σ²_total)'),
              italic=True, lang=lang)

    # FIX 3 (Alpha): removed duplicate interpretation column
    rel_rows = [
        [dim_name(r['Dimension'], lang),
         alpha_badge(r['Cronbach Alpha'], lang)]
        for r in result['reliability']
    ]
    add_styled_table(doc,
        [t('dim_lbl',   lang, 'Dimension'),
         t('alpha_lbl', lang, "Cronbach's Alpha")],
        rel_rows, lang=lang)

    poor = [r for r in result['reliability'] if r['Cronbach Alpha'] < 0.7]
    if poor:
        names = ', '.join([dim_name(r['Dimension'], lang) for r in poor])
        warn  = (f'⚠ {names}: '
                 f'{"ከ0.7 በታች ያለ አልፋ — የጥያቄ ዝርዝሮቹን ይከልሱ።" if lang == "am" else "Alpha below 0.7 — consider reviewing questionnaire items."}')
        body_text(doc, warn, color=cfg.ORANGE, lang=lang)
    add_page_break(doc)


def add_regression_section(doc, result, lang, cfg):
    """Render Section 5 — Regression Analysis with beta strength."""
    logger.info("Building regression section.")
    reg = result.get('regression')

    styled_heading(doc,
        t('s5', lang, '5. Multiple Linear Regression Analysis'), lang=lang)
    add_divider(doc)
    body_text(doc, t('s5_body', lang,
        'Multiple linear regression identifies which dimensions significantly '
        'predict Customer Satisfaction.'), lang=lang)
    body_text(doc, t('s5_model', lang,
        'SAT = β₀ + β₁·SD + β₂·TQ + β₃·PP + β₄·COM + ε'),
              italic=True, lang=lang)

    if not reg:
        body_text(doc,
            'ለሪግሬሽን ትንተና በቂ ውሂብ የለም።' if lang == 'am'
            else 'Insufficient data for regression analysis.', lang=lang)
        add_page_break(doc)
        return

    styled_heading(doc, t('s5_1', lang, '5.1 Model Summary'),
                   level=2, lang=lang)
    sig_label = (t('sig_yes', lang, 'Significant (p < 0.05)')
                 if reg['f_pvalue'] < 0.05
                 else t('sig_no', lang, 'Not significant'))
    model_rows = [
        [t('sample_n',  lang, 'Sample Size (N)'),  str(reg['n'])],
        ['R²',                                      str(reg['r_squared'])],
        [t('adj_r2',    lang, 'Adjusted R²'),       str(reg['adj_r_squared'])],
        [t('f_stat',    lang, 'F-Statistic'),       str(reg['f_statistic'])],
        [t('model_p',   lang, 'Model p-value'),     str(reg['f_pvalue'])],
        [t('model_sig', lang, 'Model Significance'), sig_label],
    ]
    add_styled_table(doc,
        ['Parameter' if lang == 'en' else 'መለኪያ',
         t('value', lang, 'Value')],
        model_rows, lang=lang)

    # FIX 7: small-N caveat
    if reg['n'] < 30:
        caveat = (
            f'ማሳሰቢያ፡ ናሙናው {reg["n"]} ብቻ ስለሆነ እና 4 ትንበያዎች ስላሉ '
            f'የስታቲስቲካዊ ጠቀሜታ ዝቅተኛ ሊሆን ይችላል። '
            f'ትክክለኛ ውጤት ለማግኘት ቢያንስ 30 ምላሾች ያስፈልጋሉ።'
            if lang == 'am' else
            f'Note: With only N={reg["n"]} observations and 4 predictors, '
            f'statistical power is low. Non-significant p-values may reflect '
            f'sample size rather than true effect absence. '
            f'A minimum of 30 respondents is recommended for reliable results.'
        )
        body_text(doc, caveat, color=cfg.ORANGE, lang=lang)

    # FIX 4 (VIF): column removed — VIF requires raw dataframe, not available here
    styled_heading(doc, t('s5_2', lang, '5.2 Coefficients'),
                   level=2, lang=lang)
    coef_rows = []
    for c in reg['coefficients']:
        strength = _beta_strength(c['Beta'], lang)
        sig_cell = (('✔ ' + t('sig_yes', lang, 'Yes'))
                    if c['Significant'] == 'Yes'
                    else t('sig_no', lang, 'No'))
        coef_rows.append([
            dim_name(c['Variable'], lang),
            c['Beta'],
            c['Std Error'],
            c['t-value'],
            c['p-value'],
            sig_cell,
            strength,
        ])
    add_styled_table(doc,
        [t('variable',     lang, 'Variable'),
         'Beta (β)',
         t('std_err',      lang, 'Std Error'),
         't-value', 'p-value',
         t('significant',  lang, 'Significant'),
         t('beta_strength',lang, 'Beta Strength')],
        coef_rows, lang=lang)

    # Regression bar chart — FIX 3: handles negative betas correctly
    chart_labels = [dim_name(c['Variable'], lang) for c in reg['coefficients']]
    chart_values = [c['Beta'] for c in reg['coefficients']]
    chart_title  = ('Regression Coefficients (β)' if lang == 'en'
                    else 'የሪግሬሽን ቅንጅቶች (β)')
    _insert_chart(doc, 'bar', chart_labels, chart_values,
                  chart_title, lang=lang)

    sig = [c for c in reg['coefficients'] if c['Significant'] == 'Yes']
    if sig:
        top = max(sig, key=lambda x: abs(x['Beta']))
        dn  = dim_name(top['Variable'], lang)
        msg = (f'{dn} (β = {top["Beta"]}): '
               f'{"በደንበኛ እርካታ ላይ ትልቁን ጠቃሚ ተጽዕኖ ያሳደረ ልኬት ነው።" if lang == "am" else "Strongest significant predictor of Customer Satisfaction (p < 0.05)."}')
        body_text(doc, msg, color=cfg.BLUE, lang=lang)
    else:
        msg = ('በ p < 0.05 ጠቃሚ ተጽዕኖ ያሳደረ ልኬት አልተገኘም — ናሙናውን ማሳደግ ያስቡ።'
               if lang == 'am' else
               'No predictor reached significance at p < 0.05 — '
               'consider increasing sample size.')
        body_text(doc, msg, color=cfg.ORANGE, lang=lang)
    add_page_break(doc)


def add_correlation_section(doc, result, lang, cfg):
    """Render Section 6 — Correlation Matrix."""
    logger.info("Building correlation section.")
    styled_heading(doc, t('s6', lang, '6. Correlation Matrix'), lang=lang)
    add_divider(doc)
    body_text(doc, t('s6_body', lang,
        'Pearson correlation coefficients between dimension scores. '
        'Values closer to 1.0 indicate strong positive relationships.'),
              lang=lang)

    corr = result.get('correlation', {})
    if corr:
        dim_names_list   = list(corr.keys())
        translated_names = [dim_name(d, lang) for d in dim_names_list]
        corr_rows = []
        for i, row_name in enumerate(dim_names_list):
            row = [translated_names[i]]
            for col_name in dim_names_list:
                val = corr[row_name].get(col_name, '')
                row.append(str(val) if val != '' else '—')
            corr_rows.append(row)
        add_styled_table(doc, [''] + translated_names, corr_rows, lang=lang)
    add_page_break(doc)


def add_nps_section(doc, result, lang, cfg):
    """Render Section 7 — NPS (only if nps data exists)."""
    nps = result.get('nps')
    if not nps:
        return
    logger.info("Building NPS section.")
    styled_heading(doc,
        t('s7_nps', lang, '7. Net Promoter Score (NPS)'), lang=lang)
    add_divider(doc)
    body_text(doc, t('s7_nps_body', lang,
        'NPS measures client loyalty. '
        'NPS = ((Promoters − Detractors) / Total) × 100'), lang=lang)

    nps_rows = [
        [t('promoters',  lang, 'Promoters (9–10)'), str(nps['promoters']),
         f"{safe_pct(nps['promoters'],  nps['total'])}%"],
        [t('passives',   lang, 'Passives (7–8)'),   str(nps['passives']),
         f"{safe_pct(nps['passives'],   nps['total'])}%"],
        [t('detractors', lang, 'Detractors (0–6)'), str(nps['detractors']),
         f"{safe_pct(nps['detractors'], nps['total'])}%"],
        [t('total', lang, 'Total'), str(nps['total']), '100%'],
    ]
    add_styled_table(doc,
        ['', t('count', lang, 'Count'), t('percentage', lang, 'Percentage')],
        nps_rows, lang=lang)

    # FIX 5: short pie chart labels — no parentheticals
    pie_labels = (['ፕሮሞተሮች', 'ፓሲቭ', 'ዲትራክተሮች'] if lang == 'am'
                  else ['Promoters', 'Passives', 'Detractors'])
    pie_values = [nps['promoters'], nps['passives'], nps['detractors']]
    if sum(pie_values) > 0:
        chart_title = ('NPS Distribution' if lang == 'en' else 'NPS ስርጭት')
        _insert_chart(doc, 'pie', pie_labels, pie_values, chart_title,
                      colors=['#2e7d32', '#1a6fa8', '#c62828'], lang=lang)

    p  = doc.add_paragraph()
    r1 = p.add_run(f'{t("nps_score", lang, "NPS Score")}: ')
    r1.bold = True
    r1.font.size = Pt(12)
    if lang == 'am': _apply_amharic_font(r1)
    r2 = p.add_run(f'{nps["nps_score"]}  —  {_nps_interp(nps["nps_score"], lang)}')
    r2.font.size = Pt(14)
    r2.bold = True
    r2.font.color.rgb = cfg.GREEN if nps['nps_score'] >= 30 else cfg.ORANGE
    if lang == 'am': _apply_amharic_font(r2)
    doc.add_paragraph()
    add_page_break(doc)


def add_insights_section(doc, result, lang, cfg):
    """
    Render the Key Insights section.
    FIX 4: Regenerates insights at render time in the correct language
    so this function is safe to call directly without going through
    build_report() first.
    """
    logger.info("Building insights section.")
    from analysis import generate_insights as _gen_insights

    # Regenerate insights in the correct language at render time
    try:
        insights = _gen_insights(
            result.get('csi', []),
            result.get('regression'),
            result.get('reliability', []),
            result.get('overall_csi', 0),
            lang=lang,
        )
    except Exception as exc:
        logger.warning("Insight regeneration failed (%s) — using pre-built.", exc)
        insights = result.get('insights', [])  # safe fallback

    nps = result.get('nps')
    sn  = 8 if nps else 7
    styled_heading(doc,
        f'{sn}. {t("insights_title", lang, "Key Insights")}', lang=lang)
    add_divider(doc)
    body_text(doc, t('insights_body', lang,
        'The following insights are derived from the analysis:'), lang=lang)
    for insight in insights:
        p   = doc.add_paragraph(style='List Number')
        run = p.add_run(insight)
        run.font.size = Pt(ReportConfig.BODY_SIZE)
        run.font.color.rgb = DARK
        if lang == 'am':
            _apply_amharic_font(run)
    doc.add_paragraph()


def add_recommendations_section(doc, result, lang, cfg):
    """Render the Recommendations section."""
    logger.info("Building recommendations section.")
    nps = result.get('nps')
    sn  = 9 if nps else 8
    styled_heading(doc,
        f'{sn}. {t("recs_title", lang, "Recommendations")}', lang=lang)
    add_divider(doc)
    body_text(doc, t('recs_body', lang,
        'Recommendations derived directly from analysis findings:'), lang=lang)
    for rec in dynamic_recommendations(result, lang):
        p   = doc.add_paragraph(style='List Number')
        run = p.add_run(rec)
        run.font.size = Pt(ReportConfig.BODY_SIZE)
        run.font.color.rgb = DARK
        if lang == 'am':
            _apply_amharic_font(run)
    doc.add_paragraph()


def add_methodology_section(doc, result, lang, cfg):
    """
    Render the Methodology Note section.
    FIX 5: Removed duplicate hardcoded footer — _add_footer() handles it.
    """
    logger.info("Building methodology section.")
    nps = result.get('nps')
    sn  = 10 if nps else 9
    styled_heading(doc,
        f'{sn}. {t("method_title", lang, "Methodology Note")}', lang=lang)
    add_divider(doc)
    method_items = [
        (t('data_coll',    lang, 'Data Collection'),
         t('data_coll_v',  lang, 'KoboToolbox online survey platform')),
        (t('scale_m',      lang, 'Scale'),
         t('scale_v',      lang, '5-point Likert (1=Very Dissatisfied, 5=Very Satisfied)')),
        (t('missing_m',    lang, 'Missing Values'),
         t('missing_mv',   lang, 'Imputed with column mean')),
        (t('csi_f',        lang, 'CSI Formula'),
         t('csi_fv',       lang, 'CSI (%) = (Mean Score / 5) × 100')),
        (t('reliability_m',lang, 'Reliability'),
         t('reliability_v',lang, "Cronbach's Alpha")),
        (t('regression_m', lang, 'Regression'),
         t('regression_v', lang, 'OLS Multiple Linear Regression')),
        (t('corr_m',       lang, 'Correlation'),
         t('corr_v',       lang, 'Pearson correlation coefficient')),
        (t('software_m',   lang, 'Software'),
         t('software_v',   lang, 'Python (pandas, numpy, statsmodels, scipy)')),
        (t('gen_date',     lang, 'Report Generated'),
         date.today().strftime('%B %d, %Y')),
    ]
    for label, val in method_items:
        add_kv(doc, label, val, lang=lang)
    doc.add_paragraph()
    # NOTE: Footer text is handled by _add_footer() in the Word footer — not repeated here


# ── Introduction (internal) ───────────────────────────────────────────────────

def _add_introduction(doc, result, lang, cfg):
    """Render Section 1 — Introduction."""
    logger.info("Building introduction section.")
    styled_heading(doc, t('s1', lang, '1. Introduction'), lang=lang)
    add_divider(doc)
    body_text(doc, t('s1_body', lang,
        'Water Works Corporation (WWC) is committed to delivering high-quality '
        'water infrastructure services to its institutional clients. This survey '
        'was designed to assess client perceptions across key service dimensions '
        'and identify areas for continuous improvement.'), lang=lang)
    body_text(doc, t('s1_body2', lang,
        'The survey was administered through KoboToolbox using a 5-point Likert '
        'scale (1 = Very Dissatisfied, 5 = Very Satisfied) across 14 performance '
        'items in four dimensions, plus three overall satisfaction items and an '
        'optional Net Promoter Score.'), lang=lang)

    styled_heading(doc, t('s1_1', lang, '1.1 Questionnaire Structure'),
                   level=2, lang=lang)
    q_rows = [
        [t('Service Delivery',    lang, 'Service Delivery'),
         'SD1–SD4',   'Q5.8–Q5.11', '4'],
        [t('Technical Quality',   lang, 'Technical Quality'),
         'TQ1–TQ3',   'Q5.5–Q5.7',  '3'],
        [t('Project Performance', lang, 'Project Performance'),
         'PP1–PP4',   'Q5.1–Q5.4',  '4'],
        [t('Communication',       lang, 'Communication'),
         'COM1–COM3', 'Q5.12–Q5.14','3'],
        [t('Customer Satisfaction',lang,'Customer Satisfaction'),
         'SAT1–SAT3', 'Q6.1–Q6.3',  '3'],
        ['NPS', 'NPS', 'Q7', '1 (0–10)'],
    ]
    add_styled_table(doc,
        [t('dim_col',   lang, 'Dimension'),
         t('vars_col',  lang, 'Variables'),
         t('q_col',     lang, 'Items'),
         t('items_col', lang, 'Count')],
        q_rows, lang=lang)

    styled_heading(doc, t('s1_2', lang, '1.2 Sample'), level=2, lang=lang)
    add_kv(doc, t('respondents', lang, 'Total Respondents'), result['n'],
           lang=lang)
    add_kv(doc, t('scale_lbl',   lang, 'Scale'),
           '1–5 Likert | 0–10 NPS', lang=lang)
    add_kv(doc, t('missing_lbl', lang, 'Missing Values'),
           t('missing_val', lang, 'Imputed with column mean'), lang=lang)
    doc.add_paragraph()
    add_page_break(doc)


# ── Master builder ────────────────────────────────────────────────────────────

def build_report(result, lang='en'):
    """
    Build a complete bilingual Word report.

    Parameters
    ----------
    result : dict   Output of analysis.run_full_analysis()
    lang   : str    'en' (English) or 'am' (Amharic)

    Returns
    -------
    bytes   Raw .docx bytes ready to write to disk or serve over HTTP.
    """
    logger.info("Starting report build (lang=%s, n=%s).", lang, result.get('n'))
    validate_result(result)
    cfg = ReportConfig()

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    org_name = ('ውሃ ሥራዎች ኮርፖሬሽን' if lang == 'am'
                else 'Water Works Corporation')
    _add_header(doc, org_name, lang=lang)
    _add_footer(doc, lang=lang)

    add_cover_page(doc, result, lang, cfg)
    add_table_of_contents(doc, lang=lang)
    add_executive_summary(doc, result, lang, cfg)
    _add_introduction(doc, result, lang, cfg)
    add_descriptive_statistics(doc, result, lang, cfg)
    add_csi_section(doc, result, lang, cfg)
    add_reliability_section(doc, result, lang, cfg)
    add_regression_section(doc, result, lang, cfg)
    add_correlation_section(doc, result, lang, cfg)
    add_nps_section(doc, result, lang, cfg)
    add_insights_section(doc, result, lang, cfg)        # FIX 4 inside
    add_recommendations_section(doc, result, lang, cfg)
    add_methodology_section(doc, result, lang, cfg)     # FIX 5 inside

    buf = io.BytesIO()
    doc.save(buf)
    logger.info("Report build complete (lang=%s).", lang)
    return buf.getvalue()


# ── Public API ────────────────────────────────────────────────────────────────

def generate_word_report(result):
    """Generate English Word report. Returns raw bytes (application/docx)."""
    return build_report(result, lang='en')


def generate_amharic_report(result):
    """
    Generate Amharic Word report. Returns raw bytes (application/docx).
    Requires 'Nyala' or 'Ebrima' font installed on the rendering system
    for correct Ethiopic script display in charts.
    """
    return build_report(result, lang='am')